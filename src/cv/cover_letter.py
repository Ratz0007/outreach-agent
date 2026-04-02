"""Cover letter generator using Claude API.

Generates a tailored cover letter alongside the CV when an application
requires one. Uses the same tailored data from the CV pipeline.
"""

import logging
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.config import PROJECT_ROOT, Secrets

logger = logging.getLogger(__name__)

EXPORTS_DIR = PROJECT_ROOT / "exports"
EXPORTS_DIR.mkdir(exist_ok=True)

COVER_LETTER_SYSTEM_PROMPT = """You are an expert cover letter writer who creates compelling,
personalised cover letters that complement (not repeat) the resume. Your cover letters:

1. Open with a strong hook — NOT "I am writing to apply for..."
2. Show genuine knowledge of the company and role
3. Connect 2-3 key achievements to the role's requirements
4. Demonstrate cultural fit and enthusiasm
5. Close with a confident call to action

Rules:
- Keep it to 3-4 paragraphs (under 350 words total)
- Use a professional but warm tone
- Reference specific company details when available
- Never repeat bullet points verbatim from the resume
- Focus on the WHY — why this company, why this role, why now
- Include concrete numbers/achievements but frame them as stories
- Never fabricate achievements — use only what's in the candidate profile"""

COVER_LETTER_USER_PROMPT = """Write a cover letter for the following application.

## CANDIDATE
Name: {name}
Location: {location}
Current headline: {headline}

## KEY ACHIEVEMENTS (from tailored resume)
{achievements}

## TARGET ROLE
Company: {company}
Role: {role}
Location: {job_location}

## JOB DESCRIPTION (excerpt)
{jd_excerpt}

## KEYWORDS TO NATURALLY INCORPORATE
{keywords}

Write the cover letter body ONLY (no addresses, no date, no "Dear Hiring Manager" — I'll add formatting).
Return ONLY the letter body text, 3-4 paragraphs, under 350 words."""


def _generate_cover_letter_text(
    profile: dict,
    job,
    tailored_data: dict,
    region_fmt: dict,
) -> str | None:
    """Generate cover letter text using Claude API."""
    if not Secrets.ANTHROPIC_API_KEY:
        logger.warning("No ANTHROPIC_API_KEY — skipping cover letter generation")
        return None

    try:
        import anthropic
        client = anthropic.Anthropic(api_key=Secrets.ANTHROPIC_API_KEY)

        # Extract top achievements from tailored bullets
        achievements = []
        for exp in tailored_data.get("experience", []):
            for bullet in exp.get("bullets", [])[:2]:
                achievements.append(f"- {bullet}")
        achievements_text = "\n".join(achievements[:6])

        keywords = tailored_data.get("keywords", [])

        prompt = COVER_LETTER_USER_PROMPT.format(
            name=profile.get("name", "Ratin Sharma"),
            location=profile.get("location", "Dublin, Ireland"),
            headline=profile.get("headline", ""),
            achievements=achievements_text,
            company=job.company,
            role=job.role,
            job_location=job.location or "Not specified",
            jd_excerpt=(job.description or "")[:2000],
            keywords=", ".join(keywords[:10]),
        )

        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=1000,
            system=COVER_LETTER_SYSTEM_PROMPT,
            messages=[{"role": "user", "content": prompt}],
        )

        text = response.content[0].text.strip()
        logger.info(f"Cover letter generated ({len(text)} chars)")
        return text

    except Exception as e:
        logger.error(f"Cover letter generation failed: {e}")
        return None


def _build_cover_letter_doc(
    name: str,
    email: str,
    phone: str,
    location: str,
    company: str,
    role: str,
    body_text: str,
    font_name: str = "Calibri",
) -> str:
    """Build the cover letter as a .docx file."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(11)

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Sender info (top right)
    sender_para = doc.add_paragraph()
    sender_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sender_lines = [name, location]
    if email:
        sender_lines.append(email)
    if phone:
        sender_lines.append(phone)
    sender_run = sender_para.add_run("\n".join(sender_lines))
    sender_run.font.name = font_name
    sender_run.font.size = Pt(10)
    sender_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    sender_para.paragraph_format.space_after = Pt(12)

    # Date
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(date.today().strftime("%B %d, %Y"))
    date_run.font.name = font_name
    date_run.font.size = Pt(11)
    date_para.paragraph_format.space_after = Pt(12)

    # Greeting
    greeting_para = doc.add_paragraph()
    greeting_run = greeting_para.add_run("Dear Hiring Manager,")
    greeting_run.font.name = font_name
    greeting_run.font.size = Pt(11)
    greeting_para.paragraph_format.space_after = Pt(8)

    # Body paragraphs
    for paragraph in body_text.split("\n\n"):
        paragraph = paragraph.strip()
        if paragraph:
            body_para = doc.add_paragraph()
            body_run = body_para.add_run(paragraph)
            body_run.font.name = font_name
            body_run.font.size = Pt(11)
            body_para.paragraph_format.space_after = Pt(8)

    # Closing
    closing_para = doc.add_paragraph()
    closing_para.paragraph_format.space_before = Pt(4)
    closing_run = closing_para.add_run("Best regards,")
    closing_run.font.name = font_name
    closing_run.font.size = Pt(11)
    closing_para.paragraph_format.space_after = Pt(4)

    name_para = doc.add_paragraph()
    name_run = name_para.add_run(name)
    name_run.font.name = font_name
    name_run.font.size = Pt(11)
    name_run.bold = True

    # Save
    safe_company = _sanitize_filename(company)
    safe_role = _sanitize_filename(role)
    today = date.today().strftime("%Y%m%d")
    filename = f"Cover_Letter_{safe_company}_{safe_role}_{today}.docx"
    filepath = EXPORTS_DIR / filename

    doc.save(str(filepath))
    logger.info(f"Cover letter saved: {filepath}")
    return str(filepath)


def _sanitize_filename(text: str) -> str:
    """Remove unsafe characters from filename."""
    import re
    return re.sub(r'[^\w\-_. ]', '', text).strip().replace(' ', '_')


def generate_cover_letter(
    profile: dict,
    job,
    tailored_data: dict,
    region_fmt: dict,
) -> str | None:
    """Generate a cover letter for a job application.

    Args:
        profile: Master profile dict
        job: JobShortlist ORM object
        tailored_data: Tailored CV data from Claude (with keywords, experience, etc.)
        region_fmt: Region format dict

    Returns:
        File path of generated cover letter, or None if generation fails
    """
    # Generate the text
    body_text = _generate_cover_letter_text(profile, job, tailored_data, region_fmt)

    if not body_text:
        return None

    # Build the document
    font_name = region_fmt.get("font", "Calibri")

    filepath = _build_cover_letter_doc(
        name=profile.get("name", "Ratin Sharma"),
        email=profile.get("email", "ratinsharma99@gmail.com"),
        phone=profile.get("phone", ""),
        location=profile.get("location", "Dublin, Ireland"),
        company=job.company,
        role=job.role,
        body_text=body_text,
        font_name=font_name,
    )

    return filepath
