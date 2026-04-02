"""Region-aware python-docx CV builder -> PDF — Stage 6.

ATS-optimized CV generation rules:
  - NO images, NO tables, NO graphics, NO icons
  - Standard fonts (region-specific): Calibri, Arial
  - Standard headings using region-specific labels
  - Keywords from JD appear naturally throughout
  - Section order follows regional conventions
"""

import logging
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.config import PROJECT_ROOT

logger = logging.getLogger(__name__)

EXPORTS_DIR = PROJECT_ROOT / "exports"
EXPORTS_DIR.mkdir(exist_ok=True)

# Dark navy for headings
HEADING_COLOR = RGBColor(0x1A, 0x1A, 0x2E)
SUBTEXT_COLOR = RGBColor(0x55, 0x55, 0x55)
LIGHT_TEXT = RGBColor(0x66, 0x66, 0x66)


def _sanitize_filename(text: str) -> str:
    """Remove unsafe characters from filename."""
    return re.sub(r'[^\w\-_. ]', '', text).strip().replace(' ', '_')


def _set_font(run, font_name: str, size: int, bold: bool = False,
              italic: bool = False, color: RGBColor = None):
    """Apply consistent font styling to a run."""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def _add_section_heading(doc: Document, text: str, font_name: str):
    """Add a section heading (level 2) with consistent formatting."""
    heading = doc.add_heading(text.upper(), level=2)
    for run in heading.runs:
        run.font.name = font_name
        run.font.color.rgb = HEADING_COLOR


def _add_paragraph(doc: Document, text: str, font_name: str = "Calibri",
                   size: int = 11, bold: bool = False, italic: bool = False,
                   color: RGBColor = None, space_after: int = 4,
                   alignment=None):
    """Add a paragraph with consistent formatting."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_font(run, font_name, size, bold, italic, color)
    para.paragraph_format.space_after = Pt(space_after)
    if alignment:
        para.alignment = alignment
    return para


def _add_bullet(doc: Document, text: str, font_name: str = "Calibri", size: int = 11):
    """Add a bullet point."""
    para = doc.add_paragraph(style="List Bullet")
    para.clear()
    run = para.add_run(text)
    _set_font(run, font_name, size)
    para.paragraph_format.space_after = Pt(2)
    return para


def _build_header(doc: Document, name: str, contact_parts: list[str],
                  font_name: str, region_fmt: dict):
    """Build the name + contact info header."""
    # Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(name)
    _set_font(name_run, font_name, 18, bold=True, color=HEADING_COLOR)
    name_para.paragraph_format.space_after = Pt(2)

    # Contact line
    contact_text = "  |  ".join(p for p in contact_parts if p)
    if contact_text:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run(contact_text)
        _set_font(contact_run, font_name, 10, color=SUBTEXT_COLOR)
        contact_para.paragraph_format.space_after = Pt(8)


def _build_summary_section(doc: Document, summary: str, label: str, font_name: str):
    """Build the professional summary section."""
    _add_section_heading(doc, label, font_name)
    _add_paragraph(doc, summary, font_name=font_name, size=11, space_after=8)


def _build_experience_section(doc: Document, experience: list[dict],
                              label: str, font_name: str, font_size: int,
                              bullet_style: str):
    """Build the experience section with all roles and bullets."""
    _add_section_heading(doc, label, font_name)

    for exp in experience:
        # Title — Company line
        title_text = exp.get("title", "")
        company_text = exp.get("company", "")
        title_line = f"{title_text}  —  {company_text}" if company_text else title_text

        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title_line)
        _set_font(title_run, font_name, font_size, bold=True)
        title_para.paragraph_format.space_after = Pt(1)

        # Location + dates on same line
        meta_parts = []
        if exp.get("location"):
            meta_parts.append(exp["location"])
        if exp.get("dates"):
            meta_parts.append(exp["dates"])

        if meta_parts:
            meta_para = doc.add_paragraph()
            meta_run = meta_para.add_run("  |  ".join(meta_parts))
            _set_font(meta_run, font_name, 10, italic=True, color=LIGHT_TEXT)
            meta_para.paragraph_format.space_after = Pt(2)

        # Bullets
        for bullet in exp.get("bullets", []):
            _add_bullet(doc, bullet, font_name, font_size)

        # Spacer between roles
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(4)


def _build_education_section(doc: Document, education: list[dict],
                             label: str, font_name: str, font_size: int):
    """Build the education section."""
    _add_section_heading(doc, label, font_name)

    for edu in education:
        degree = edu.get("degree", "")
        school = edu.get("school", "")
        dates = edu.get("dates", "")

        edu_line = f"{degree}  —  {school}"
        if dates:
            edu_line += f"  ({dates})"

        edu_para = doc.add_paragraph()
        edu_run = edu_para.add_run(edu_line)
        _set_font(edu_run, font_name, font_size, bold=True)
        edu_para.paragraph_format.space_after = Pt(2)

        if edu.get("thesis"):
            thesis_para = doc.add_paragraph()
            thesis_run = thesis_para.add_run(f"Thesis: {edu['thesis']}")
            _set_font(thesis_run, font_name, 10, italic=True)
            thesis_para.paragraph_format.space_after = Pt(4)


def _build_skills_section(doc: Document, skills, label: str,
                          font_name: str, font_size: int,
                          skills_format: str):
    """Build skills section based on region format.

    skills_format:
      - 'categorized': grouped by category with headers
      - 'detailed': with proficiency levels
      - 'flat': simple comma-separated list
    """
    _add_section_heading(doc, label, font_name)

    if isinstance(skills, dict):
        if skills_format == "flat":
            # Flatten all skills into one list
            all_skills = []
            for category, items in skills.items():
                if isinstance(items, list):
                    all_skills.extend(items)
            skills_text = "  |  ".join(all_skills)
            _add_paragraph(doc, skills_text, font_name=font_name, size=font_size)

        elif skills_format == "detailed":
            # Show with category headers and proficiency
            for category, items in skills.items():
                if isinstance(items, list) and items:
                    cat_para = doc.add_paragraph()
                    cat_run = cat_para.add_run(f"{category.replace('_', ' ').title()}: ")
                    _set_font(cat_run, font_name, font_size, bold=True)
                    items_run = cat_para.add_run(", ".join(items))
                    _set_font(items_run, font_name, font_size)
                    cat_para.paragraph_format.space_after = Pt(3)

        else:  # categorized (default)
            for category, items in skills.items():
                if isinstance(items, list) and items:
                    cat_para = doc.add_paragraph()
                    cat_run = cat_para.add_run(f"{category.replace('_', ' ').title()}: ")
                    _set_font(cat_run, font_name, font_size, bold=True)
                    items_run = cat_para.add_run(", ".join(items))
                    _set_font(items_run, font_name, font_size)
                    cat_para.paragraph_format.space_after = Pt(3)

    elif isinstance(skills, list):
        skills_text = "  |  ".join(skills)
        _add_paragraph(doc, skills_text, font_name=font_name, size=font_size)


# ── Section builders keyed by name ───────────────────────────────────
SECTION_BUILDERS = {
    "summary": "_build_summary",
    "experience": "_build_experience",
    "education": "_build_education",
    "skills": "_build_skills",
}


def build_cv(
    name: str,
    email: str,
    linkedin: str,
    location: str,
    summary: str,
    experience: list[dict],
    education: list[dict],
    skills,
    company: str,
    role: str,
    region_fmt: dict = None,
    phone: str = "",
) -> str:
    """Build an ATS-optimised CV using region-specific formatting.

    Args:
        name: Full name
        email: Email address
        linkedin: LinkedIn URL
        location: Location
        summary: Professional summary (tailored)
        experience: List of {company, title, location, dates, bullets}
        education: List of {degree, school, dates, thesis?}
        skills: Dict of {category: [skills]} or list of skills
        company: Target company (for filename)
        role: Target role (for filename)
        region_fmt: Region format dict from regions.py (optional)
        phone: Phone number (optional, included for some regions)

    Returns:
        Full file path of the generated .docx
    """
    # Default to UK/Ireland format if no region specified
    if region_fmt is None:
        from src.cv.regions import get_region_format
        region_fmt = get_region_format("")

    font_name = region_fmt.get("font", "Calibri")
    font_size = region_fmt.get("font_size", 11)
    margins = region_fmt.get("margins_inches", 0.7)
    section_order = region_fmt.get("section_order", ["summary", "experience", "skills", "education"])
    bullet_style = region_fmt.get("bullet_style", "achievement")
    skills_format = region_fmt.get("skills_format", "categorized")

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(font_size)

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(margins)
        section.right_margin = Inches(margins)

    # ── Header ───────────────────────────────────────────────────
    contact_parts = [location]
    if phone and (region_fmt.get("include_address", False) or region_fmt.get("include_dob", False)):
        # Regions that include more personal details also show phone
        contact_parts.append(phone)
    contact_parts.append(email)
    contact_parts.append(linkedin)

    _build_header(doc, name, contact_parts, font_name, region_fmt)

    # ── Build sections in region-specific order ──────────────────
    for section_name in section_order:
        if section_name == "summary":
            label = region_fmt.get("summary_label", "Professional Summary")
            _build_summary_section(doc, summary, label, font_name)

        elif section_name == "experience":
            label = region_fmt.get("experience_label", "Experience")
            _build_experience_section(
                doc, experience, label, font_name, font_size, bullet_style
            )

        elif section_name == "education":
            label = region_fmt.get("education_label", "Education")
            _build_education_section(
                doc, education, label, font_name, font_size
            )

        elif section_name == "skills":
            label = region_fmt.get("skills_label", "Skills")
            _build_skills_section(
                doc, skills, label, font_name, font_size, skills_format
            )

    # ── Save ─────────────────────────────────────────────────────
    safe_company = _sanitize_filename(company)
    safe_role = _sanitize_filename(role)
    today = date.today().strftime("%Y%m%d")
    filename = f"Ratin_Sharma_{safe_company}_{safe_role}_{today}.docx"
    filepath = EXPORTS_DIR / filename

    doc.save(str(filepath))
    logger.info(f"CV saved: {filepath} (region: {region_fmt.get('name', 'default')})")

    return str(filepath)
